import os
from modules.Return_Latest_File import ret_latest_file
import docx

def writeOnFile(title, content, filetype = 'docx'):
    
    name = 'NewContent'
    pathNewContent = './data/Rewritten Blog/'
    
    if os.path.exists(pathNewContent):
        filename = f'{pathNewContent}/{name}_{str(ret_latest_file(name)[1])}'
    else:
        os.mkdir(pathNewContent)
        filename = f'{pathNewContent}/{name}_{str(ret_latest_file(name)[1])}'
    
    # Writing Text File.
    if filetype == 'txt':
        
        f = open(f'{filename}.{filetype}', 'w')

        f.write(f'{title}\n')

        f.write(content)
        f.close()

    # Writing Docx File.
    if filetype == 'docx':        
        doc = docx.Document()

        doc.add_heading(title)

        doc.add_paragraph(content)
        
        # now save the document to a location 
        doc.save(f'{filename}.{filetype}')

writeOnFile("Head", "dd")